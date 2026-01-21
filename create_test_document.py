# create_test_document.py
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_comprehensive_test_document():
    doc = Document()
    
    # ===== 1. CORRECTE BRANDING =====
    # Titel in Calibri (correct font)
    title = doc.add_heading('Technical Report: Laméco AI Style Checker Test', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 2. FONT ISSUES =====
    # Sectie met verschillende fonts
    doc.add_heading('1. Font Compliance Testing', level=1)
    
    # Correct font (Calibri) - default
    p1 = doc.add_paragraph('This text is in the correct font (Calibri).')
    
    # Font issue: Arial
    p2 = doc.add_paragraph()
    run2 = p2.add_run('ERROR: This text is in Arial font instead of Calibri.')
    run2.font.name = 'Arial'
    
    # Font issue: Times New Roman  
    p3 = doc.add_paragraph()
    run3 = p3.add_run('ERROR: This text is in Times New Roman.')
    run3.font.name = 'Times New Roman'
    
    # Font weight issue: Bold body text
    p4 = doc.add_paragraph()
    run4 = p4.add_run('ERROR: Body text should not be bold. This is bold text.')
    run4.bold = True
    
    # ===== 3. COLOR ISSUES =====
    doc.add_heading('2. Color Compliance Testing', level=1)
    
    # Correct color (brand color #1A3CE0)
    p5 = doc.add_paragraph('This text is in the correct brand color.')
    run5 = p5.runs[0]
    run5.font.color.rgb = RGBColor(0x1A, 0x3C, 0xE0)  # #1A3CE0
    
    # Color issue: Red text
    p6 = doc.add_paragraph()
    run6 = p6.add_run('ERROR: This text is red (#FF0000) instead of brand blue.')
    run6.font.color.rgb = RGBColor(255, 0, 0)  # Red
    
    # Color issue: Green text
    p7 = doc.add_paragraph()
    run7 = p7.add_run('ERROR: This text is green (#00FF00).')
    run7.font.color.rgb = RGBColor(0, 255, 0)  # Green
    
    # ===== 4. TONE OF VOICE TESTING =====
    doc.add_heading('3. Tone of Voice Testing', level=1)
    
    # Formal tone (should be OK or marked as formal)
    p8 = doc.add_paragraph('Geachte heer, wij verzoeken u vriendelijk om dit document te beoordelen. Ondergetekende verklaart dat alle gegevens correct zijn.')
    
    # Informal tone (should be flagged)
    p9 = doc.add_paragraph('Hey! Kun je dit ff checken? Super bedankt hoor! Doei!')
    
    # Neutral tone
    p10 = doc.add_paragraph('Het document wordt morgen opgeleverd. De deadline is 17:00 uur.')
    
    # Mixed tone
    p11 = doc.add_paragraph('Beste collega, graag willen we je vragen om dit NA TE KIJKEN! Dankjewel!')
    
    # ===== 5. GRAMMAR & SPELLING =====
    doc.add_heading('4. Grammar & Spelling Testing', level=1)
    
    # Dutch grammar errors
    p12 = doc.add_paragraph('Hun hebben het document gemaakt. Dit is fout want het moet "zij" zijn.')
    
    # Informal abbreviations
    p13 = doc.add_paragraph('Kun je dit mss ff nakijken? Iig voor morgen graag.')
    
    # Spelling mistakes
    p14 = doc.add_paragraph('Dit dokument bevat spelfouten en gramatika fouten.')
    
    # ===== 6. TECHNICAL CONTENT (jouw document) =====
    doc.add_heading('5. Technical Report Content', level=1)
    
    # Voeg jouw technische content toe
    tech_content = """
Technical Report: Assignment 1 – Electric Concepts

Introduction
This report documents the practical experiment conducted to analyze electric concepts, specifically focusing on series and parallel resistor circuits.

Procedure
1. Step 1: Circuit Setup
   A circuit was constructed using a breadboard, an Arduino, and two resistors.

2. Step 2: Resistor Measurement
   • R1 = 4.7K Ω ±5%
   • R2 = 10 KΩ ±5%

3. Step 3: Parallel Circuit Configuration
   The circuit was reconfigured to a parallel arrangement.

4. Step 4: Verification of OHM LAW
   I = U/R
   • R1 = 5 / 4.7 KΩ = 1.064 A
   • R2 = 5 / 10 KΩ = 0.0005 A

5. Step 5: Equivalent Resistance Calculation
   1 / R-equivalent = (1 / R1) + (1 / R2)
   (1 / 10000) + (1 / 1000) = 1/0.0011 = 909 Ω

6. Step 6: Current Calculation
   I = V / R-equivalent = 5V / 6609 Ω = 0.756 mA

7. Step 7: Power Dissipation
   P = V x I = 5 x 0.000756 = 3.78 mW

Conclusions
The experiment successfully demonstrated the principles of electric circuits.
"""
    
    doc.add_paragraph(tech_content)
    
    # ===== 7. HEADINGS WITH CORRECT FORMAT =====
    doc.add_heading('6. Correct Headings Format', level=2)
    doc.add_paragraph('This is a level 2 heading with proper formatting.')
    
    doc.add_heading('7. Another Section', level=3)
    doc.add_paragraph('Level 3 heading - should be detected correctly.')
    
    # ===== 8. TABLE WITH ISSUES =====
    doc.add_heading('8. Table Formatting', level=1)
    
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Component'
    header_cells[1].text = 'Value'
    header_cells[2].text = 'Tolerance'
    
    # Data rows
    row1 = table.rows[1].cells
    row1[0].text = 'R1'
    row1[1].text = '4.7KΩ'
    row1[2].text = '5%'
    
    row2 = table.rows[2].cells
    row2[0].text = 'R2'
    row2[1].text = '10KΩ'
    row2[2].text = '5%'
    
    # ===== 9. BULLET POINTS =====
    doc.add_heading('9. List Examples', level=1)
    
    doc.add_paragraph('Key findings:', style='List Bullet')
    doc.add_paragraph('Resistors within tolerance', style='List Bullet')
    doc.add_paragraph('Circuit operates safely', style='List Bullet')
    doc.add_paragraph('Measurements accurate', style='List Bullet')
    
    # ===== 10. MIXED LANGUAGE (Dutch/English) =====
    doc.add_heading('10. Language Mix Test', level=1)
    
    p15 = doc.add_paragraph('This sentence is in English. ')
    run15 = p15.add_run('Deze zin is in het Nederlands. ')
    run15.font.color.rgb = RGBColor(255, 0, 0)  # Red for visibility
    p15.add_run('Back to English for the conclusion.')
    
    # ===== 11. SPECIAL CHARACTERS =====
    doc.add_heading('11. Special Characters & Symbols', level=1)
    doc.add_paragraph('Resistance: 4.7KΩ ±5% → Current: 1.064mA @ 5V')
    doc.add_paragraph('Formulas: P = V × I, R = ρ × L/A')
    doc.add_paragraph('Temperature range: -40°C to +85°C')
    
    # Save document
    doc.save('test_document_all_features.docx')
    print('Test document created: test_document_all_features.docx')
    print('\nThis document tests:')
    print('1. ✓ Font compliance (Calibri vs Arial/Times New Roman)')
    print('2. ✓ Color compliance (brand blue vs red/green)')
    print('3. ✓ Tone of voice (formal, informal, neutral)')
    print('4. ✓ Grammar & spelling (Dutch errors, informal abbreviations)')
    print('5. ✓ Technical content (your electric concepts report)')
    print('6. ✓ Headings structure')
    print('7. ✓ Tables')
    print('8. ✓ Lists')
    print('9. ✓ Mixed language')
    print('10. ✓ Special characters')

if __name__ == '__main__':
    create_comprehensive_test_document()